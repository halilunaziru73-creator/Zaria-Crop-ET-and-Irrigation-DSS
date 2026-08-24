"""
report_docx.py
---------------
Assembles a full Word (.docx) report for a farm owner: every number in it comes from
the same results/season/downstream/today_multi dicts the GUI already computed for the
entered crop/temperature/humidity/area — nothing here is recalculated differently or
guessed. XAI (plain-language) explanations come from xai_explain.py, which templates
real values into pre-written sentences, not a separate model.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from . import xai_explain as xai
from . import farm_map as fm
from . import irrigation_types as it
from . import config as cfg

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")
FIG_DIR = os.path.join(OUT_DIR, "figures")

ACCENT_RGB = RGBColor(0x2B, 0x7A, 0x78)
DARK_RGB = RGBColor(0x17, 0x25, 0x2A)


class _FigureCounter:
    def __init__(self):
        self.n = 0

    def next(self):
        self.n += 1
        return self.n


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT_RGB if level > 1 else DARK_RGB
    return h


def _kv_table(doc, rows, col_widths=(2.5, 3.5)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
        row[0].paragraphs[0].runs[0].font.bold = True
    for row in table.rows:
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
    return table


def _add_figure(doc, path, caption, counter: "_FigureCounter", width_in=6.0):
    """caption: a plain descriptive string (no manual 'Figure N' prefix, no dash) --
    the sequential 'Figure N:' label is generated here from the shared counter, so
    every figure in the document is numbered in true document order, starting at 1."""
    if not path or not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        n = counter.next()
        cap = doc.add_paragraph(f"Figure {n}: {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def _build_header_footer(doc, farm_owner_name: str):
    """Adds a repeating page header (the app's logo) and footer (author credit + photo)
    to every page of the document. Explicit header/footer distances are set so the
    content can never overlap with the page body regardless of viewer/renderer."""
    from . import logo as logo_mod
    section = doc.sections[0]
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        logo_path = logo_mod.generate_logo()
        run = hp.add_run()
        run.add_picture(logo_path, height=Inches(0.3))
        text_run = hp.add_run("  Zaria Crop ET and Irrigation DSS")
        text_run.bold = True
        text_run.font.size = Pt(10)
        text_run.font.name = "Calibri"
    except Exception:
        text_run = hp.add_run("Zaria Crop ET and Irrigation DSS")
        text_run.bold = True
        text_run.font.size = Pt(10)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    photo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets",
                               "author_naziru_halilu.png")
    try:
        if os.path.exists(photo_path):
            run = fp.add_run()
            run.add_picture(photo_path, height=Inches(0.18))
            text_run = fp.add_run("  Author: Naziru Halilu")
            text_run.italic = True
            text_run.font.size = Pt(8)
        else:
            text_run = fp.add_run("Author: Naziru Halilu")
            text_run.italic = True
            text_run.font.size = Pt(8)
    except Exception:
        text_run = fp.add_run("Author: Naziru Halilu")
        text_run.italic = True
        text_run.font.size = Pt(8)


def generate_full_report(results: dict, downstream: dict, today_multi: dict, season: dict,
                          model, farm_owner_name: str, area_ha: float,
                          farm_lat: float = None, farm_lon: float = None,
                          water_flow_direction_deg: float = 200.0,
                          output_dir: str = OUT_DIR, output_path: str = None) -> str:
    crop = results["crop"]
    ti = results["temperature_input"]

    # Clear any report-specific figures left over from a PREVIOUS farm owner/session
    # in the same run (process_diagram_<name>.png, farm_map_<name>.png, etc.) --
    # otherwise files from different owners accumulate side by side and can show up
    # together as apparent "duplicate" figures in the app's Figures tab.
    stale_prefixes = ("process_diagram_", "farm_map_", "growth_filmstrip_",
                       "growth_simulation_", "monitoring_chart_", "qgis_terrain_",
                       "qgis_comparison_", "panel_et_climate_", "panel_soil_efficiency_")
    if os.path.isdir(FIG_DIR):
        for fname in os.listdir(FIG_DIR):
            if fname.startswith(stale_prefixes):
                try:
                    os.remove(os.path.join(FIG_DIR, fname))
                except OSError:
                    pass

    # farm map + irrigation recommendation (schematic — see farm_map.py for the
    # disclosed limitation: no GEE/internet access in this environment)
    net_irrig = results["water"]["net_irrigation_mm"]
    gross_irrig = results["water"]["gross_irrigation_mm"]
    recommendation = fm.recommend_irrigation_method(crop["key"], net_irrig, gross_irrig)
    lat = farm_lat if farm_lat is not None else results["site"]["latitude_deg"]
    lon = farm_lon if farm_lon is not None else 7.65  # Zaria approx longitude if not supplied
    farm_map_path = fm.generate_farm_map(
        farm_owner_name or "Farm", lat, lon, area_ha, crop["display_name"], recommendation,
        crop_key=crop["key"],
        fname=f"farm_map_{(farm_owner_name or 'farm').replace(' ', '_')}.png")

    # QGIS-style terrain study-area layout, sized to this farm's real area, unique
    # per farm (see qgis_layout.py for the disclosed regional-model limitation)
    from . import qgis_layout as ql
    growing_season_etc = results["et"].get("growing_season_etc_mm")
    local_range = crop.get("local_etc_range_mm")
    qgis_result = ql.generate_zaria_study_area_layout(
        farm_owner_name or "Farm", area_ha=area_ha, lat=lat, lon=lon,
        crop_display_name=crop["display_name"], growing_season_etc_mm=growing_season_etc,
        local_etc_range_mm=local_range)

    # Productivity estimate computed here (early) so it's available for the Executive
    # Summary as well as Section 12 later, without recomputing.
    from . import productivity as prodmod
    prod_est = prodmod.estimate_productivity(crop["key"], downstream["dwb"], area_ha)

    doc = Document()
    fig_counter = _FigureCounter()
    _build_header_footer(doc, farm_owner_name)

    # ---------------- Title page ----------------
    title = doc.add_heading("Zaria Crop ET and Irrigation DSS: Farm Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Prepared for: {farm_owner_name or '(farm owner not specified)'}")
    run.font.size = Pt(14)
    run.font.bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Crop: {crop['display_name']}   |   Farm area: {area_ha} ha   |   "
        f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).font.size = Pt(10)
    _add_figure(doc, farm_map_path, "Farm layout showing the delineated boundary, planted crop, and "
                                     "irrigation infrastructure.", fig_counter)

    # ---------------- How to Use This Report ----------------
    doc.add_page_break()
    _heading(doc, "How to Use This Report")
    doc.add_paragraph(
        "This report is organised so that a busy reader can act on it quickly, while still giving a "
        "technical reviewer full traceability back to the underlying data and equations for every number."
    )
    guide_items = [
        ("Executive Summary", "Start here for the numbers that matter for a same-day irrigation "
         "decision: predicted water use, net irrigation required, recommended method, and estimated yield."),
        ("Sections 1-2 (Overview, Dashboard)", "The live inputs (temperature, humidity) and how they "
         "translate into today's crop water use, with a full plain-text data dump for anyone who wants "
         "every intermediate number."),
        ("Section 3-4 (Soil Water, Irrigation Schedule)", "How much water the root zone can hold, and "
         "the practical, stage-specific schedule to follow this season."),
        ("Section 5 (Efficiency & Water Budget)", "Where water is lost between the source and the crop, "
         "and the full seasonal water balance, verified to close to zero."),
        ("Sections 6-7 (Irrigation Method)", "Which delivery method (drip, sprinkler, furrow, "
         "basin/flood) fits this farm and crop, and how much water each would need."),
        ("Sections 8-11 (Figures, Terrain, Growth Simulation)", "Supporting charts, this farm's own "
         "boundary and terrain, and a stage-by-stage field simulation from germination to harvest."),
        ("Section 12 (Productivity)", "An estimated yield range, specific to this run's own computed "
         "water-stress trajectory, not a generic figure."),
        ("Section 13-14 (Monitoring, Weekly Balance)", "What to check routinely for maximum "
         "productivity, and a week-by-week ledger of the whole season."),
        ("Sections 15-16 (Methodology, Glossary)", "Full transparency on every equation, data source, "
         "and term used, for a technical reviewer or a second opinion."),
    ]
    table_guide = doc.add_table(rows=1, cols=2)
    table_guide.style = "Light Grid Accent 1"
    hdr = table_guide.rows[0].cells
    hdr[0].text = "Section"; hdr[1].text = "What it's for"
    hdr[0].paragraphs[0].runs[0].font.bold = True
    hdr[1].paragraphs[0].runs[0].font.bold = True
    for name, desc in guide_items:
        row = table_guide.add_row().cells
        row[0].text = name
        row[1].text = desc

    # ---------------- Executive Summary ----------------
    doc.add_page_break()
    _heading(doc, "Executive Summary")
    sched_summary = results["schedule"]
    rec_summary = results["recommended_schedule"]
    wb_summary = results["water_budget"]
    doc.add_paragraph(
        f"This report covers {area_ha} ha of {crop['display_name']} at {farm_owner_name or 'this farm'}, "
        f"located near Zaria, Kaduna State (Lat {lat:.4f}, Lon {lon:.4f}). It was generated on "
        f"{datetime.now().strftime('%Y-%m-%d at %H:%M')} using an entered temperature of "
        f"{ti['temperature_c']}\u00b0C and humidity of {ti['humidity_pct']}%."
    )
    _kv_table(doc, [
        ("Predicted crop water use today (ETc)", f"{results['et']['today_predicted_etc']} mm/day"),
        ("Growing-season crop water use", f"{results['et'].get('growing_season_etc_mm')} mm over "
                                           f"{results['et'].get('growing_season_days')} days"),
        ("Net irrigation required this season", f"{wb_summary.get('net_irrigation_mm')} mm"),
        ("Recommended irrigation method", recommendation["recommended_method"]),
        ("Number of recommended irrigation events", str(rec_summary.get("n_events"))),
        ("Critical water-sensitivity window", rec_summary.get("critical_window", "n/a")),
        ("Overall irrigation system efficiency", f"{results['efficiency']['Ep']}%"),
        ("Estimated yield", f"{prod_est['estimated_yield_t_ha']} t/ha "
                             f"({prod_est['estimated_total_production_t']} t total)"),
        ("Water balance check", f"closes to {wb_summary.get('balance_residual_mm')} mm residual"),
    ])
    doc.add_paragraph(
        "Full detail, methodology, and every supporting figure and table for each number above follow "
        "in the numbered sections below."
    )

    # ---------------- Site and Climate Normals ----------------
    doc.add_page_break()
    _heading(doc, "Site and Climate Normals")
    doc.add_paragraph(
        f"Monthly climate normals for {results['site']['town']}, {results['site']['state']} "
        f"(Lat {results['site']['latitude_deg']}\u00b0N, {results['site']['elevation_m']} m a.s.l.), "
        f"averaged from a 28-year daily weather archive (2000-2027). These are the same climatology "
        f"values driving every day-of-year lookup elsewhere in this report (wind, solar radiation, and "
        f"rainfall on days other than today, which uses your entered temperature and humidity directly)."
    )
    import calendar
    month_rows = []
    for m in range(1, 13):
        doys_in_month = [d for d in range(1, 366)
                          if __import__("datetime").date(2021, 1, 1).toordinal() + d - 1 <=
                          __import__("datetime").date(2021, 12, 31).toordinal()
                          and __import__("datetime").date.fromordinal(
                              __import__("datetime").date(2021, 1, 1).toordinal() + d - 1).month == m]
        rows_for_month = [model.doy_climatology.get(str(d)) for d in doys_in_month]
        rows_for_month = [r for r in rows_for_month if r]
        if not rows_for_month:
            continue
        tmax_m = round(sum(r["tmax"] for r in rows_for_month) / len(rows_for_month), 1)
        tmin_m = round(sum(r["tmin"] for r in rows_for_month) / len(rows_for_month), 1)
        wind_m = round(sum(r["wind"] for r in rows_for_month) / len(rows_for_month), 2)
        solar_m = round(sum(r["solar"] for r in rows_for_month) / len(rows_for_month), 1)
        rain_m = round(sum(r["rainfall"] for r in rows_for_month), 1)
        month_rows.append((calendar.month_name[m], tmax_m, tmin_m, wind_m, solar_m, rain_m))

    table_clim = doc.add_table(rows=1, cols=6)
    table_clim.style = "Light Grid Accent 1"
    hdr = table_clim.rows[0].cells
    for i, h in enumerate(["Month", "Mean Tmax (\u00b0C)", "Mean Tmin (\u00b0C)", "Mean Wind (m/s)",
                           "Mean Solar (MJ/m\u00b2/day)", "Total Rainfall (mm)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for mn, tmax_m, tmin_m, wind_m, solar_m, rain_m in month_rows:
        row = table_clim.add_row().cells
        row[0].text = mn
        row[1].text = str(tmax_m)
        row[2].text = str(tmin_m)
        row[3].text = str(wind_m)
        row[4].text = str(solar_m)
        row[5].text = str(rain_m)

    # ---------------- 1. Overview ----------------
    _heading(doc, "1. Overview")
    doc.add_paragraph(xai.explain_overview(results))
    _kv_table(doc, [
        ("Date", ti["date"]), ("Input temperature", f"{ti['temperature_c']} \u00b0C"),
        ("Input humidity", f"{ti['humidity_pct']} %"), ("Crop coefficient (Kc) today", ti["kc_today"]),
        ("Season status", crop["season_label"]),
        ("Predicted ETc today", f"{results['et']['today_predicted_etc']} mm/day"),
        ("Growing-season ETc", f"{results['et'].get('growing_season_etc_mm')} mm "
                                f"over {results['et'].get('growing_season_days')} days"),
        ("Locally-reported ETc range", f"{crop['local_etc_range_mm'][0]}-{crop['local_etc_range_mm'][1]} mm"),
    ])

    from . import process_diagram as pdmod
    proc_values = pdmod.compute_process_values(ti["temperature_c"], ti["humidity_pct"], today_multi,
                                                downstream, downstream["taw_raw"],
                                                today_index=season.get("dap_today"))
    proc_values["temp_c"] = ti["temperature_c"]
    proc_values["rh_pct"] = ti["humidity_pct"]
    proc_diagram_path = pdmod.plot_process_diagram(
        proc_values, farm_owner_name or "This Farm", crop["display_name"],
        fname=f"process_diagram_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    doc.add_paragraph(
        "The diagram below shows exactly how today's entered temperature and humidity flow through "
        "this pipeline's dual crop-coefficient (Kcb/Ke) engine, the dynamic root-zone depth, and the "
        "soil-water depletion check, to decide whether the crop is running at its full potential water "
        "use rate or under a stress reduction today. Every value shown is this run's own, not a "
        "generic illustration."
    )
    _add_figure(doc, proc_diagram_path, "ET / soil-water decision process for today's inputs.",
                fig_counter, width_in=4.2)

    # ---------------- 2. Dashboard ----------------
    _heading(doc, "2. Dashboard Summary")
    from . import report as rpt
    dash = doc.add_paragraph(rpt.build_dashboard(results))
    for run in dash.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)

    _heading(doc, "2.1 ET Methods Used", level=2)
    doc.add_paragraph(
        "Each reference-ET method evaluated for today's entered temperature and humidity, with its own "
        "result alongside a brief description. The predicted ETc used throughout this report is the "
        "ensemble mean of whichever methods have sufficient data to run."
    )
    method_descriptions = [
        ("FAO-56 Penman-Monteith", "The internationally standard method; combines energy balance and "
         "aerodynamic terms, requires temperature, humidity, wind, and solar radiation."),
        ("ASCE Standardized Penman-Monteith", "A US-standardized variant of Penman-Monteith with "
         "fixed surface resistance coefficients for a short reference crop."),
        ("Original/Modified Penman", "An earlier combination-equation formulation, predating the "
         "FAO-56 standardisation."),
        ("Priestley-Taylor", "A simplified radiation-driven method, using an empirical multiplier "
         "on the energy-balance term alone."),
        ("Hargreaves-Samani", "A temperature-only method, useful when humidity/wind/radiation data "
         "are unavailable; used as the trend model backbone for this pipeline's thermal-unit "
         "regression."),
        ("Makkink", "A radiation-based method calibrated for humid climates, requiring only "
         "temperature and solar radiation."),
        ("Turc", "A simplified radiation and temperature method, developed for humid conditions."),
        ("Thornthwaite", "A temperature-and-daylength-only method, originally developed for climate "
         "classification rather than irrigation engineering."),
        ("Blaney-Criddle", "An older, temperature-and-daylength empirical method still referenced in "
         "irrigation engineering practice."),
        ("Dalton-type Mass Transfer", "An aerodynamic-only method based on vapour pressure deficit "
         "and wind speed, without an energy-balance term."),
    ]
    table_methods = doc.add_table(rows=1, cols=4)
    table_methods.style = "Light Grid Accent 1"
    hdr = table_methods.rows[0].cells
    for i, h in enumerate(["Method", "Description", "Today's ET0 (mm/day)", "Today's ETc (mm/day)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    method_results = today_multi.get("methods", {})
    for name, desc in method_descriptions:
        row = table_methods.add_row().cells
        row[0].text = name
        row[1].text = desc
        result = method_results.get(name, {})
        if result.get("status") == "OK":
            row[2].text = str(result.get("et0_mm_day", "n/a"))
            row[3].text = str(result.get("etc_mm_day", "n/a"))
        else:
            row[2].text = "insufficient data"
            row[3].text = "insufficient data"

    et0_values = [r["et0_mm_day"] for r in method_results.values() if r.get("status") == "OK"]
    if len(et0_values) >= 2:
        spread = round(max(et0_values) - min(et0_values), 2)
        note = doc.add_paragraph()
        note.add_run("Why the results differ between methods: ").bold = True
        note.add_run(
            f"today's ET0 estimates span {round(min(et0_values), 2)} to {round(max(et0_values), 2)} mm/day "
            f"(a spread of {spread} mm/day) because each method uses a different subset of the weather "
            f"drivers. Full energy-balance methods (FAO-56/ASCE Penman-Monteith, Priestley-Taylor) weigh "
            f"solar radiation and wind alongside temperature and humidity; temperature-only methods "
            f"(Hargreaves-Samani, Thornthwaite, Blaney-Criddle) ignore radiation and wind entirely, which "
            f"pulls them away from the energy-balance methods whenever conditions are unusually windy, "
            f"cloudy, or clear for the season; and the aerodynamic-only method (Dalton-type) responds "
            f"only to vapour pressure deficit and wind, missing the radiation-driven component entirely. "
            f"None of these methods is \"wrong\" -- the spread itself is a genuine measure of how much a "
            f"single day's ET0 depends on which weather variables are actually available, which is why "
            f"this pipeline reports the ensemble mean rather than any one method alone."
        )

    # ---------------- 3. Soil Water ----------------
    _heading(doc, "3. Soil Water")
    doc.add_paragraph(xai.explain_soil_water(results))
    sw = results["soil_water"]
    _kv_table(doc, [
        ("Field capacity", f"{sw['field_capacity_pct']} %"),
        ("Permanent wilting point", f"{sw['pwp_pct']} %"),
        ("Root zone depth", f"{sw['root_zone_depth_m']} m"),
        ("Maximum allowable depletion (MAD)", sw["MAD"]),
        ("Total Available Water (TAW)", f"{sw['TAW_mm']} mm"),
        ("Readily Available Water (RAW)", f"{sw['RAW_mm']} mm"),
    ])
    from . import soil_profile as spmod
    dwb_today_idx = season.get("dap_today", -1)
    dwb_list = downstream["dwb"]
    today_dwb_rec = (dwb_list[dwb_today_idx] if dwb_list and -len(dwb_list) <= dwb_today_idx < len(dwb_list)
                      else None)
    zr_series_data = downstream.get("zr_series")
    zr_today_val = (zr_series_data[dwb_today_idx]
                     if zr_series_data and -len(zr_series_data) <= dwb_today_idx < len(zr_series_data) else None)
    kc_today_val = ti["kc_today"]
    maturity_val = max(0.05, min(1.0, kc_today_val / 1.2)) if kc_today_val else 0.5
    soil_profile_path = spmod.plot_soil_profile(
        crop["key"], crop["display_name"], farm_owner_name or "This Farm",
        cfg.DEFAULT_SOIL, downstream["taw_raw"],
        storage_mm=today_dwb_rec.storage_mm if today_dwb_rec else None,
        depletion_mm=today_dwb_rec.depletion_mm if today_dwb_rec else None,
        root_zone_depth_m=zr_today_val, maturity=maturity_val,
        fname=f"soil_profile_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    _add_figure(doc, soil_profile_path, "Soil-water profile for this crop and farm, with the crop shown "
                                          "at its current stage above a labelled soil column.", fig_counter)

    # ---------------- 4. Irrigation Schedule ----------------
    _heading(doc, "4. Irrigation Schedule")
    rec_sched = results["recommended_schedule"]
    si = rec_sched["stage_intervals_days"]
    if rec_sched["n_events"] > 0:
        crit = doc.add_paragraph()
        crit.add_run(f"\u26a0 Critical water-sensitivity window: ").bold = True
        crit.add_run(rec_sched["critical_window"])
        doc.add_paragraph(
            f"The irrigation interval VARIES by growth stage: nursery/germination needs frequent light "
            f"watering, while the flowering/fruiting stage above is this crop's tightest, most critical "
            f"interval of the whole season. Depth per application also varies, following the day's own "
            f"root-zone capacity as the crop's roots deepen."
        )
        _kv_table(doc, [
            ("Nursery/Germination interval", f"every {si['Initial']} days"),
            ("Vegetative interval", f"every {si['Development']} days"),
            ("Flowering/Fruiting interval", f"every {si['Mid-season']} days"),
            ("Maturity interval", f"every {si['Late-season']} days"),
        ])
        mm_to_m3 = area_ha * 10
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Day", "Stage", "Net (mm)", "Gross (mm)", "Net Vol (m\u00b3)", "Gross Vol (m\u00b3)"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for e in rec_sched["events"]:
            row = table.add_row().cells
            row[0].text = str(e["day"])
            row[1].text = str(e["stage"])
            row[2].text = str(e["net_irrigation_mm"])
            row[3].text = str(e["gross_irrigation_mm"])
            row[4].text = f"{e['net_irrigation_mm'] * mm_to_m3:,.1f}"
            row[5].text = f"{e['gross_irrigation_mm'] * mm_to_m3:,.1f}"
    else:
        doc.add_paragraph(
            "No irrigation is recommended for the entered conditions, rainfall/soil moisture is "
            "expected to meet crop demand across the growing season."
        )
    note_diff = doc.add_paragraph()
    note_diff.add_run(
        "Note: the practical schedule above (a fixed, stage-specific interval a farmer can follow) is "
        "distinct from the day-by-day simulated water balance reported in Section 5 (which triggers "
        "irrigation exactly when depletion crosses the safe threshold, giving irregular real-world "
        "timing). The two totals will not exactly match; the schedule above is a planning "
        "simplification of the same underlying physics, not a different set of numbers."
    ).italic = True

    # ---------------- 5. Efficiency & Water Budget ----------------
    _heading(doc, "5. Efficiency & Water Budget")
    doc.add_paragraph(xai.explain_efficiency_and_budget(results))
    src1 = doc.add_paragraph()
    src1.add_run("Source: standard irrigation-engineering assumptions for the configured system "
                 "(furrow, Ea/Ec/Ed defaults), not measured for this specific farm.").italic = True
    eff = results["efficiency"]
    _kv_table(doc, [
        ("Conveyance efficiency (Ec)", f"{eff['Ec']} %"),
        ("Distribution efficiency (Ed)", f"{eff['Ed']} %"),
        ("Application efficiency (Ea)", f"{eff['Ea']} %"),
        ("Overall system efficiency (Ep)", f"{eff['Ep']} %"),
    ])
    src2 = doc.add_paragraph()
    src2.add_run("Source: predicted from the entered temperature/humidity via the trained weather/ET "
                 "model, combined with the system efficiency assumptions above.").italic = True
    wb = results["water_budget"]
    _kv_table(doc, [(k.replace("_", " ").title(), f"{v} mm") for k, v in wb.items() if v is not None])

    # ---------------- 6. Irrigation Method Recommendation ----------------
    _heading(doc, "6. Irrigation Method Recommendation")
    doc.add_paragraph(xai.explain_irrigation_recommendation(recommendation))

    # ---------------- 7. Water Required by Irrigation Method (this farm's area) ----------------
    _heading(doc, "7. Water Required by Irrigation Method: This Farm's Area")
    method_rows = it.water_required_by_method(net_irrig, area_ha)
    doc.add_paragraph(
        f"For {farm_owner_name or 'this farm'}'s {area_ha} ha, delivering the computed net irrigation "
        f"requirement of {net_irrig} mm to the root zone requires the following gross water volumes "
        f"depending on the irrigation method chosen. Lower-efficiency methods need more water applied "
        f"at the field to deliver the same amount to the crop."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Method", "Application Eff. (%)", "Gross Depth (mm)", "Gross Volume (m\u00b3)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for r in method_rows:
        row = table.add_row().cells
        row[0].text = r["method"]
        row[1].text = str(r["application_efficiency_pct"])
        row[2].text = str(r["gross_depth_mm"])
        row[3].text = f"{r['gross_volume_m3']:,.1f}"
    for r in method_rows:
        doc.add_paragraph(f"\u2022 {r['method']}: {r['note']}")

    # ---------------- 8. Additional Figures ----------------
    _heading(doc, "8. Additional Figures")
    from . import panel_layout as pl
    panel1_path = pl.assemble_grid_panel([
        (os.path.join(FIG_DIR, "thermal_unit_regression.png"), "Thermal-unit (GDD) regression, pooled training data"),
        (os.path.join(FIG_DIR, "seasonal_et_vs_reference.png"), "Crop ET vs reference ET across the season"),
        (os.path.join(FIG_DIR, "today_method_comparison.png"), "ETc today across all applicable methods"),
        (os.path.join(FIG_DIR, "cumulative_etc.png"), "Cumulative crop ET across the season"),
    ], "ET and Climate Analysis", f"panel_et_climate_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    if panel1_path:
        _add_figure(doc, panel1_path, "ET and climate analysis.", fig_counter, width_in=6.3)

    panel2_path = pl.assemble_grid_panel([
        (os.path.join(FIG_DIR, "soil_depletion.png"), "Root-zone depletion vs RAW threshold"),
        (os.path.join(FIG_DIR, "rainfall_vs_etc.png"), "Rainfall vs crop ET across the season"),
        (os.path.join(FIG_DIR, "efficiency_breakdown.png"), "Irrigation system efficiency breakdown"),
        (os.path.join(FIG_DIR, "water_budget.png"), "Seasonal water budget"),
    ], "Soil-Water Balance and System Efficiency",
       f"panel_soil_efficiency_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    if panel2_path:
        _add_figure(doc, panel2_path, "Soil-water balance and system efficiency.", fig_counter, width_in=6.3)

    # ---------------- 9. Study Area: Terrain Characterisation ----------------
    _heading(doc, "9. Study Area: Terrain Characterisation")
    fstats = qgis_result["farm_stats"]
    doc.add_paragraph(
        f"The figure below follows a standard QGIS-print-layout terrain-characterisation format "
        f"(location, altitude distribution, slope, topographic profiles, aspect), sized to this "
        f"farm's real entered area ({area_ha} ha) with an organic boundary shaped by the underlying "
        f"terrain. This farm's modelled terrain averages "
        f"{fstats['mean_elevation_m']} m a.s.l. with a mean slope of {fstats['mean_slope_pct']}%."
    )
    _add_figure(doc, qgis_result["terrain_layout"], "Terrain characterisation for this farm's own boundary.",
                fig_counter)
    if qgis_result.get("farm_comparison"):
        _add_figure(doc, qgis_result["farm_comparison"],
                    "This farm's growing-season ETc compared with the locally-reported Zaria range for this crop.",
                    fig_counter)

    # ---------------- 10. Crop Growth Simulation: Germination to Harvest ----------------
    _heading(doc, "10. Crop Growth Simulation: Germination to Harvest")
    from . import growth_simulation as gsim
    growth = gsim.plot_growth_simulation(crop["key"], crop["display_name"], farm_owner_name or "This Farm",
                                          model=model, fname=f"growth_simulation_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    doc.add_paragraph(
        f"Simulated planting (germination) date: {growth['planting_date']}. Estimated harvest date: "
        f"{growth['harvest_date']} ({growth['season_length_days']} days total). The growth-stage "
        f"boundaries and crop coefficient (Kc) values are the same ones driving the ET/irrigation "
        f"calculations elsewhere in this report; the canopy-cover curve shown is a proxy derived "
        f"directly from Kc (not a remote-sensed or field-measured canopy observation)."
    )
    _add_figure(doc, growth["path"], "Growth-stage timeline and canopy-cover proxy, germination to harvest.", fig_counter)

    # ---------------- 11. Stage-by-Stage Field Simulation & Water Accounting ----------------
    _heading(doc, "11. Stage-by-Stage Field Simulation and Water Accounting")
    from . import growth_video as gv
    stage_rows = gv.compute_stage_water_balance(crop["key"], area_ha, model, downstream)
    filmstrip_path = gv.generate_filmstrip(crop["key"], crop["display_name"], farm_owner_name or "This Farm",
                                            stage_rows, fname=f"growth_filmstrip_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    doc.add_paragraph(
        "A field picture for each growth stage (procedurally illustrated, not a photograph, no "
        "internet access to source real field imagery in this environment), alongside that stage's "
        "own water accounting: crop coefficient, ET split into evaporation (soil) vs transpiration "
        "(plant), mean soil-moisture retained, and irrigation required for this farm's area. Evaporation/"
        "transpiration are split using the standard FAO-56 canopy-cover proxy, not a lysimeter measurement."
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Stage", "Days", "Kc", "ETc (mm)", "Evap. (mm)", "Transp. (mm)", "Irrigation (mm)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for r in stage_rows:
        row = table.add_row().cells
        row[0].text = r["label"]
        row[1].text = f"{r['day_range'][0]}-{r['day_range'][1]}"
        row[2].text = str(r["kc_mean"])
        row[3].text = str(r["etc_stage_mm"])
        row[4].text = str(r["evaporation_mm"])
        row[5].text = str(r["transpiration_mm"])
        row[6].text = f"{r['net_irrigation_mm']} ({r['net_irrigation_m3']:,.0f} m\u00b3)"
    _add_figure(doc, filmstrip_path, "Field-stage simulation with per-stage water accounting.", fig_counter)
    gif_note = doc.add_paragraph()
    gif_note.add_run(
        "An animated version of this simulation (a lightweight GIF cycling through the stages) is also "
        "available in the app's Figures tab / outputs folder; Word documents cannot embed animated images."
    ).italic = True

    # ---------------- 12. Estimated Productivity (Water-Yield Response Model) ----------------
    doc.add_page_break()
    _heading(doc, "12. Estimated Productivity: Water-Yield Response Model")
    doc.add_paragraph(
        f"This estimate is specific to {farm_owner_name or 'this farm'}'s own simulated water-stress "
        f"trajectory for the entered conditions, not a generic figure. A Water Stress Index (WSI) of "
        f"{prod_est['water_stress_index']} was computed from this run's own daily soil-water-balance "
        f"records (0 = never stressed, 1 = permanently at the depletion limit). A small neural network "
        f"(trained on the published FAO-33 Doorenbos & Kassam water-yield response relationship for "
        f"this crop's yield response factor, Ky = {prod_est['ky_yield_response_factor']}) converts that "
        f"stress index into a relative yield of {prod_est['relative_yield_pct']}% of potential."
    )
    _kv_table(doc, [
        ("Water Stress Index (this run)", prod_est["water_stress_index"]),
        ("Yield response factor (Ky)", prod_est["ky_yield_response_factor"]),
        ("Relative yield (Ya/Ym)", f"{prod_est['relative_yield_pct']}%"),
        ("Reference potential yield range (Ym)", f"{prod_est['potential_yield_range_t_ha'][0]}-{prod_est['potential_yield_range_t_ha'][1]} t/ha"),
        ("Estimated yield, this farm", f"{prod_est['estimated_yield_t_ha']} t/ha "
                                        f"(range {prod_est['estimated_yield_range_t_ha'][0]}-{prod_est['estimated_yield_range_t_ha'][1]})"),
        ("Estimated total production", f"{prod_est['estimated_total_production_t']} t over {area_ha} ha"),
    ])
    note = doc.add_paragraph()
    note.add_run(
        "The potential-yield range (Ym) is a literature reference for improved-practice irrigated "
        "production in the wider West African/Nigerian context, not a measurement specific to this "
        "farm's soil or variety. Treat the estimated yield as an order-of-magnitude planning figure, "
        "and replace Ym with local trial or extension-service data where available for a tighter estimate."
    ).italic = True

    # ---------------- 13. What to Monitor for Maximum Productivity ----------------
    doc.add_page_break()
    _heading(doc, "13. What to Monitor for Maximum Productivity")
    from . import monitoring as monmod
    checklist = monmod.build_monitoring_checklist(crop["key"], overall_efficiency_pct=eff["Ep"])
    doc.add_paragraph(
        f"Tailored to {crop['display_name']} and to this run's own results (e.g. system efficiency is "
        f"flagged first if it is unusually low)."
    )
    monitoring_chart_path = monmod.plot_monitoring_chart(
        crop["key"], crop["display_name"], farm_owner_name or "This Farm", overall_efficiency_pct=eff["Ep"],
        fname=f"monitoring_chart_{(farm_owner_name or 'farm').replace(' ', '_')}.png")
    _add_figure(doc, monitoring_chart_path, "Recommended monitoring frequency by topic.", fig_counter)
    for title, desc in checklist:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ---------------- 14. Weekly Water Balance Summary ----------------
    doc.add_page_break()
    _heading(doc, "14. Weekly Water Balance Summary")
    doc.add_paragraph(
        f"The full growing season, summarised week by week from the same daily soil-water-balance "
        f"simulation driving every other number in this report: crop water use (ETc), rainfall, "
        f"irrigation applied, and the root-zone depletion at the end of each week."
    )
    dwb = downstream["dwb"]
    season_days = results["et"].get("growing_season_days") or len(dwb)
    table_wk = doc.add_table(rows=1, cols=6)
    table_wk.style = "Light Grid Accent 1"
    hdr = table_wk.rows[0].cells
    for i, h in enumerate(["Week", "Days", "ETc (mm)", "Rainfall (mm)", "Irrigation (mm)", "End Depletion (mm)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    week_no = 1
    i = 0
    while i < min(season_days, len(dwb)):
        chunk = dwb[i:i + 7]
        if not chunk:
            break
        etc_sum = round(sum(d.etc_mm for d in chunk), 1)
        rain_sum = round(sum(d.rainfall_mm for d in chunk), 1)
        irr_sum = round(sum(d.irrigation_mm for d in chunk), 1)
        end_depletion = round(chunk[-1].depletion_mm, 1)
        row = table_wk.add_row().cells
        row[0].text = str(week_no)
        row[1].text = f"{i + 1}-{i + len(chunk)}"
        row[2].text = str(etc_sum)
        row[3].text = str(rain_sum)
        row[4].text = str(irr_sum)
        row[5].text = str(end_depletion)
        week_no += 1
        i += 7

    # ---------------- 15. Methodology and Data Sources ----------------
    doc.add_page_break()
    _heading(doc, "15. Methodology and Data Sources")
    doc.add_paragraph(
        "Every model, equation, and data source used to produce this report, stated explicitly."
    )
    methodology_items = [
        ("Reference evapotranspiration (ET0)", "Ten published methods are evaluated in parallel for "
         "today's entered temperature and humidity, including FAO-56 Penman-Monteith, ASCE Standardized "
         "Penman-Monteith, the original Penman equation, Priestley-Taylor, Hargreaves-Samani, Makkink, "
         "Turc, Thornthwaite, Blaney-Criddle, and Dalton-type mass transfer. The predicted ETc shown "
         "throughout this report is the ensemble mean of whichever methods have sufficient data to run "
         "for the entered conditions."),
        ("Crop coefficient (Kc)", "For maize, the real field-measured Kc calendar from a 28-year Zaria "
         "daily weather/water-balance dataset. For rice, sorghum, pepper, and cowpea, the published "
         "FAO-56 stage-wise Kc curve, rescaled to this crop's own locally-reported growing-season length "
         "and calibrated so the modelled seasonal total matches the midpoint of the locally-reported "
         "ETc range for that crop in the Zaria area."),
        ("Soil-water balance", "A daily root-zone water balance (FAO-56 Chapter 8 approach) tracking "
         "rainfall infiltration, deep percolation, crop water use, and irrigation triggers, with a "
         "root zone depth that grows from a shallow nursery-stage value toward the crop's mature depth "
         "as canopy cover develops, rather than a single fixed depth for the whole season."),
        ("Irrigation scheduling", "A standardised, stage-specific interval schedule (nursery, "
         "vegetative, flowering/fruiting, maturity), following a field-reported reference table of "
         "typical intervals for each stage and crop, with application depth computed from that "
         "day's own root-zone capacity."),
        ("Irrigation system efficiency", "Standard published conveyance, distribution, and field "
         "application efficiency ranges for the configured irrigation method (furrow by default)."),
        ("Productivity estimate", "The FAO Irrigation and Drainage Paper No. 33 (Doorenbos and Kassam, "
         "1979) water-yield response relationship, with the crop's own published yield response factor "
         "(Ky), evaluated via a small neural network trained on that published relationship across "
         "this run's own computed water-stress index. Absolute yield uses a literature-reported "
         "potential-yield range for improved-practice irrigated production in the region."),
        ("Terrain and farm layout", "This environment has no internet access, so no real satellite "
         "imagery, Google Earth Engine data, or surveyed Digital Elevation Model could be retrieved for "
         "any specific farm. The terrain and boundary figures in this report are a regional model "
         "calibrated to Zaria's documented elevation statistics, deterministic and unique per farm "
         "name and coordinates, clearly labelled on the figure itself. Replace with real surveyed data "
         "when available for field use."),
    ]
    for title, desc in methodology_items:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ---------------- 16. Glossary ----------------
    doc.add_page_break()
    _heading(doc, "16. Glossary")
    glossary = [
        ("ET0 (Reference evapotranspiration)", "The water use rate of a hypothetical, well-watered "
         "reference grass surface, driven only by weather (mm/day)."),
        ("ETc (Crop evapotranspiration)", "The actual water use rate of the specific crop grown, "
         "equal to ET0 multiplied by the crop coefficient (Kc)."),
        ("Kc (Crop coefficient)", "A multiplier, specific to the crop and its growth stage, that "
         "converts reference ET0 into crop-specific ETc."),
        ("TAW (Total Available Water)", "The total amount of water the root zone can hold between "
         "field capacity and the permanent wilting point (mm)."),
        ("RAW (Readily Available Water)", "The portion of TAW the crop can safely use before "
         "irrigation should be triggered, without yield-affecting stress (mm)."),
        ("Depletion (Dr)", "How much of the RAW/TAW has actually been used up on a given day (mm)."),
        ("VPD (Vapour Pressure Deficit)", "The difference between the saturation and actual vapour "
         "pressure of the air, driven by temperature and humidity, a direct measure of atmospheric "
         "drying power (kPa)."),
        ("Ks (Water stress coefficient)", "A 0-1 multiplier applied to potential ETc when the crop is "
         "under water stress, where 1.0 means no stress and lower values mean reduced water use "
         "due to insufficient soil moisture."),
        ("Ky (Yield response factor)", "A published, crop-specific constant describing how sensitive "
         "final yield is to a given amount of water stress over the season."),
        ("WSI (Water Stress Index)", "This report's own summary measure of how close to the safe "
         "depletion limit the crop's root zone typically was, averaged across the season (0 = never "
         "stressed, 1 = permanently at the limit)."),
        ("Ea, Ec, Ed (Irrigation efficiencies)", "Application, conveyance, and distribution "
         "efficiency: the fraction of water that survives each stage of delivery from source to the "
         "root zone without being lost."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)

    # ---------------- 17. Assumptions and Limitations ----------------
    doc.add_page_break()
    _heading(doc, "17. Assumptions and Limitations")
    doc.add_paragraph(
        "A summary, in one place, of every assumption and known limitation behind the numbers in this "
        "report, so it can be read alongside the results rather than only being scattered through the "
        "methodology notes above."
    )
    limitations = [
        ("Weather", "Today's temperature and humidity are exactly what was entered; all other days in "
         "the season use a 28-year day-of-year climatological average, not a real forecast."),
        ("Soil", "Field capacity, wilting point, and root-zone depth are standard published values for "
         "the assumed soil texture, not a laboratory analysis of this specific field."),
        ("Irrigation system", "Conveyance, distribution, and application efficiencies use standard "
         "published ranges for the selected method, not a measurement of this farm's actual "
         "infrastructure condition."),
        ("Terrain and boundary", "The terrain figures use a regional elevation model calibrated to "
         "Zaria's documented statistics rather than a survey of this specific field, clearly marked "
         "on the figure itself."),
        ("Productivity estimate", "The yield figure is an order-of-magnitude planning estimate from a "
         "published water-yield response relationship and a literature-reported potential-yield range, "
         "not a field trial result for this farm's soil or variety."),
        ("Crop calibration", "Only maize uses a field-measured crop-coefficient calendar from a 28-year "
         "Zaria dataset; the other four crops use published FAO-56 curves calibrated to locally-reported "
         "seasonal totals."),
    ]
    for title, desc in limitations:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    if output_path:
        out_path = output_path
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    else:
        os.makedirs(output_dir, exist_ok=True)
        safe_name = (farm_owner_name or "farm_report").strip().replace(" ", "_") or "farm_report"
        out_path = os.path.join(output_dir, f"{safe_name}_irrigation_report.docx")
    doc.save(out_path)
    return out_path
